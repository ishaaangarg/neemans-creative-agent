"""
Neeman's Creative Strategy Agent
=================================
Paste a product URL -> Get 10 static + 5 video ad concepts, brand-aligned and ready to produce.

Built with Streamlit + Claude API.
"""

import streamlit as st
import requests
import json
import re
import smtplib
import base64
import concurrent.futures
from datetime import datetime
from io import BytesIO
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from anthropic import Anthropic
from together import Together

# ──────────────────────────────────────────────
# PAGE CONFIG
# ──────────────────────────────────────────────
st.set_page_config(
    page_title="Neeman's Creative Strategy Agent",
    page_icon="👟",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ──────────────────────────────────────────────
# CUSTOM CSS — Neeman's brand look
# ──────────────────────────────────────────────
st.markdown(
    """
<style>
@import url('https://fonts.googleapis.com/css2?family=Bebas+Neue&family=Inter:wght@400;500;600;700&display=swap');

/* Global */
.stApp { background-color: #F3F2EE; }

/* Header */
.main-header {
    font-family: 'Bebas Neue', sans-serif;
    color: #3A3A3A;
    font-size: 2.6rem;
    text-transform: uppercase;
    letter-spacing: 3px;
    margin-bottom: 0;
}
.sub-header {
    font-family: 'Inter', sans-serif;
    color: #6B6B6B;
    font-size: 1.05rem;
    margin-top: -8px;
    margin-bottom: 24px;
}

/* Product card */
.product-card {
    background: #ffffff;
    border-radius: 14px;
    padding: 28px;
    box-shadow: 0 2px 16px rgba(0,0,0,0.06);
    margin: 12px 0 24px 0;
}

/* Concept cards */
.concept-card {
    background: #ffffff;
    border-radius: 10px;
    padding: 22px 24px;
    margin: 14px 0;
    border-left: 5px solid #AEC3AA;
    box-shadow: 0 1px 6px rgba(0,0,0,0.05);
}
.concept-card.video { border-left-color: #3A3A3A; }
.concept-card.priority { border-left-color: #D4A574; }

/* Tags */
.brand-tag {
    display: inline-block;
    background: #AEC3AA;
    color: #3A3A3A;
    padding: 3px 14px;
    border-radius: 20px;
    font-size: 0.82rem;
    margin: 3px 4px 3px 0;
    font-family: 'Inter', sans-serif;
}
.brand-tag.price {
    background: #3A3A3A;
    color: #F3F2EE;
    font-weight: 600;
}

/* Stats row */
.stat-box {
    background: #3A3A3A;
    color: #F3F2EE;
    border-radius: 10px;
    padding: 18px;
    text-align: center;
    font-family: 'Inter', sans-serif;
}
.stat-box .stat-value {
    font-family: 'Bebas Neue', sans-serif;
    font-size: 2rem;
    letter-spacing: 1px;
}
.stat-box .stat-label {
    font-size: 0.78rem;
    text-transform: uppercase;
    letter-spacing: 1px;
    opacity: 0.7;
}

/* Sidebar */
section[data-testid="stSidebar"] {
    background-color: #3A3A3A;
}
section[data-testid="stSidebar"] .stMarkdown { color: #F3F2EE; }
section[data-testid="stSidebar"] h3 { color: #EDD8C3; font-family: 'Inter', sans-serif; }
section[data-testid="stSidebar"] label { color: #F3F2EE !important; }

/* Hide Streamlit branding */
#MainMenu { visibility: hidden; }
footer { visibility: hidden; }

/* Dividers */
hr { border-color: #D4D4D0; }
</style>
""",
    unsafe_allow_html=True,
)


# ──────────────────────────────────────────────
# UTILITY FUNCTIONS
# ──────────────────────────────────────────────

def extract_handle(url: str) -> str | None:
    """Pull the product handle from a Neeman's URL."""
    m = re.search(r"/products/([a-zA-Z0-9_-]+)", url)
    return m.group(1) if m else None


def scrape_product(url: str) -> tuple[dict | None, str | None]:
    """Fetch product JSON from Neeman's Shopify API."""
    handle = extract_handle(url)
    if not handle:
        return None, "Invalid URL — could not extract product handle."
    try:
        r = requests.get(
            f"https://neemans.com/products/{handle}.json",
            timeout=15,
            headers={"User-Agent": "NeemansCreativeAgent/1.0"},
        )
        if r.status_code == 200:
            return r.json().get("product", {}), None
        return None, f"Product not found (HTTP {r.status_code})."
    except Exception as e:
        return None, f"Network error: {e}"


def product_summary(p: dict) -> dict:
    """Turn raw Shopify product JSON into a display-friendly dict."""
    variants = p.get("variants", [])
    prices = [float(v["price"]) for v in variants if v.get("price")]
    compare = [
        float(v["compare_at_price"])
        for v in variants
        if v.get("compare_at_price")
    ]
    tags = (
        p.get("tags", "").split(", ")
        if isinstance(p.get("tags"), str)
        else p.get("tags", [])
    )
    sizes = sorted(
        {v.get("option1", "") for v in variants if v.get("available")},
        key=lambda s: float(s) if s.replace(".", "").isdigit() else 0,
    )
    return {
        "title": p.get("title", ""),
        "type": p.get("product_type", ""),
        "price": f"\u20b9{min(prices):,.0f}" if prices else "N/A",
        "compare": f"\u20b9{min(compare):,.0f}" if compare else None,
        "tags": tags,
        "images": [i["src"] for i in p.get("images", [])[:6]],
        "sizes": sizes,
        "body": re.sub(r"<[^>]+>", " ", p.get("body_html", "")),
        "handle": p.get("handle", ""),
        "published": p.get("published_at", "")[:10],
    }


def load_brand_context() -> str | None:
    """Load the brand knowledge-base markdown."""
    import os

    for name in ("brand_context.md", "neemans_brand.md"):
        for base in (
            os.path.dirname(os.path.abspath(__file__)),
            os.getcwd(),
            os.path.join(os.path.dirname(os.path.abspath(__file__)), ".."),
        ):
            path = os.path.join(base, name)
            if os.path.isfile(path):
                with open(path, encoding="utf-8") as f:
                    return f.read()
    return None


SYSTEM_PROMPT_TEMPLATE = """You are the **Neeman's Creative Strategy Agent** — an elite D2C performance-creative strategist who lives and breathes the Neeman's brand.

─── BRAND KNOWLEDGE BASE ───
{brand_context}

─── YOUR MANDATE ───
When given scraped product data you MUST execute this pipeline:

**STEP 1 — PRODUCT ANALYSIS**
Analyze price position, materials, use-cases, audience fit, whether to lead with comfort / fashion / price / sustainability.

**STEP 2 — BRAND LENS**
Apply every relevant guideline:
• Personality → "The Conscious Explorer"
• Tone → Talk, don't write · Ask questions · Avoid clichés · US English
• 3 Pillars → Comfort (rational), Fashion (social), Responsibility (personal)
• Typography → Bebas Neue Pro ALL-CAPS headlines, DIN Regular body
• Color → Coal #3A3A3A · Ecru #F3F2EE · Wool #EDD8C3 · Forest #AEC3AA · Others per manual
• Rule → No price-first hooks for premium / full-MRP products
• Rule → Sustainability is layered, never the lead

**STEP 3 — GENERATE 10 STATIC AD CONCEPTS**
Each concept MUST include these fields in a markdown table:
| Field | Detail |
|---|---|
| concept_title | … |
| hook_text | ≤ 12 words |
| visual_direction | … |
| primary_copy | … |
| supporting_copy | … |
| ad_format | single / carousel / story |
| creative_reasoning | Why this works for THIS product |
| reference_brand | Real D2C brand that inspired the pattern |
| image_prompt | A detailed text-to-image prompt (50-80 words) describing the hero visual for this ad. Include: subject, composition, lighting, color palette, mood, camera angle, style (e.g. studio product shot, lifestyle, flat lay). Do NOT include any text/words in the image description — images should be text-free. |

**STEP 4 — GENERATE 5 VIDEO AD CONCEPTS**
Each concept MUST include:
| Field | Detail |
|---|---|
| concept_title | … |
| hook_text | ≤ 12 words |
| duration | e.g. 15s / 30s / 45s |
| scene_breakdown | 3-5 numbered scenes |
| script_notes | Key voiceover / super lines |
| performance_signals | When to use this ad (funnel stage, audience) |
| reference_brand | Real brand that inspired this |
| storyboard_prompts | Exactly 3 text-to-image prompts (one per key frame) for a visual storyboard. Each prompt is 40-60 words describing one scene. Format as: FRAME 1: prompt /// FRAME 2: prompt /// FRAME 3: prompt. No text/words in images. |

**STEP 5 — PRIORITIZE**
• Rank all 15 concepts by expected ROAS potential.
• Flag TOP 3 must-produce concepts.
• Note production effort (Low / Medium / High).
• Ensure ≥ 2 concepts carry a sustainability angle.
• Include a Creative Reference Map: which competitor brands → which concepts.

─── OUTPUT FORMAT ───
Use these exact markdown headings so the app can parse your response:

# PRODUCT SUMMARY
# BRAND LENS APPLIED
# STATIC AD CONCEPTS
## STATIC 1: "title"
… (repeat for 2-10)
# VIDEO AD CONCEPTS
## VIDEO 1: "title"
… (repeat for 2-5)
# PRIORITY MATRIX
# CREATIVE REFERENCE MAP

Be hyper-specific. A junior designer should be able to produce every ad from your descriptions alone.
"""


def build_user_prompt(product: dict, url: str) -> str:
    """Build the user message containing scraped product data."""
    tags = product.get("tags", "")
    if isinstance(tags, list):
        tags = ", ".join(tags)
    variants = product.get("variants", [])
    body = re.sub(r"<[^>]+>", " ", product.get("body_html", ""))
    body = re.sub(r"\s+", " ", body).strip()[:2500]

    lines = [
        "Generate a complete Creative Strategy for this Neeman's product:\n",
        f"**Product:** {product.get('title', '')}",
        f"**Type:** {product.get('product_type', '')}",
        f"**Handle:** {product.get('handle', '')}",
        f"**Price:** \u20b9{variants[0].get('price', 'N/A') if variants else 'N/A'}",
    ]
    if variants and variants[0].get("compare_at_price"):
        lines.append(
            f"**Compare-at Price:** \u20b9{variants[0]['compare_at_price']}"
        )
    lines += [
        f"**Tags:** {tags}",
        f"**Published:** {product.get('published_at', 'N/A')[:10]}",
        f"\n**Variants ({len(variants)}):**",
    ]
    for v in variants[:6]:
        lines.append(
            f"  - {v.get('title','')}: \u20b9{v.get('price','')} "
            f"(available={v.get('available',False)})"
        )
    lines += [
        f"\n**Product Description:**\n{body}",
        f"\n**Product URL:** {url}",
        "\n---\nNow execute the full creative strategy pipeline.",
    ]
    return "\n".join(lines)


def stream_strategy(product, brand_ctx, api_key, model, url):
    """Generator — yields text chunks from Claude API."""
    client = Anthropic(api_key=api_key)
    system = SYSTEM_PROMPT_TEMPLATE.format(brand_context=brand_ctx)
    user = build_user_prompt(product, url)
    with client.messages.stream(
        model=model,
        max_tokens=8000,
        system=system,
        messages=[{"role": "user", "content": user}],
    ) as stream:
        for text in stream.text_stream:
            yield text


def send_email(to, subject, body, cfg):
    """Send strategy via SMTP (Gmail)."""
    msg = MIMEMultipart("alternative")
    msg["Subject"] = subject
    msg["From"] = cfg["email"]
    msg["To"] = to
    msg.attach(MIMEText(body, "plain"))
    html = (
        "<div style='font-family:Arial,sans-serif;max-width:800px;margin:auto'>"
        f"<pre style='white-space:pre-wrap'>{body}</pre></div>"
    )
    msg.attach(MIMEText(html, "html"))
    try:
        with smtplib.SMTP(cfg.get("server", "smtp.gmail.com"), int(cfg.get("port", 587))) as s:
            s.starttls()
            s.login(cfg["email"], cfg["password"])
            s.send_message(msg)
        return True, None
    except Exception as e:
        return False, str(e)


def make_docx(title, strategy_md):
    """Convert strategy markdown to a DOCX file and return bytes."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()

    # Title
    h = doc.add_heading(f"Creative Strategy: {title}", level=0)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x3A, 0x3A, 0x3A)
    doc.add_paragraph(
        f"Generated {datetime.now().strftime('%B %d, %Y')} by Neeman's Creative Strategy Agent"
    )
    doc.add_paragraph("")

    for line in strategy_md.split("\n"):
        stripped = line.strip()
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("|") and "---" not in stripped:
            # Simple table row → paragraph
            cells = [c.strip() for c in stripped.split("|") if c.strip()]
            doc.add_paragraph("  |  ".join(cells), style="List Bullet")
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped:
            doc.add_paragraph(stripped)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ──────────────────────────────────────────────
# IMAGE GENERATION (Together AI)
# ──────────────────────────────────────────────

# Model configs
VISUAL_MODELS = {
    "quick": {
        "model": "black-forest-labs/FLUX.1-schnell",
        "steps": 4,
        "label": "Flux Schnell",
        "supports_image_ref": False,
    },
    "shoe_plus": {
        "model": "black-forest-labs/FLUX.1-kontext-dev",
        "steps": 28,
        "label": "Flux Kontext",
        "supports_image_ref": True,
    },
}


def generate_single_image(
    prompt: str,
    api_key: str,
    mode: str = "quick",
    width: int = 1024,
    height: int = 1024,
    ref_image_url: str | None = None,
) -> tuple[str | None, str | None]:
    """Generate one image via Together AI. Returns (base64_string, error_message)."""
    try:
        cfg = VISUAL_MODELS.get(mode, VISUAL_MODELS["quick"])
        client = Together(api_key=api_key)

        kwargs = dict(
            model=cfg["model"],
            prompt=prompt,
            width=width,
            height=height,
            steps=cfg["steps"],
            n=1,
            response_format="b64_json",
        )

        # Kontext: pass the product shoe as reference image
        if cfg["supports_image_ref"] and ref_image_url:
            kwargs["image_url"] = ref_image_url

        resp = client.images.generate(**kwargs)

        if resp.data and resp.data[0].b64_json:
            return resp.data[0].b64_json, None
        # Fallback: if URL is returned instead, download it
        if resp.data and hasattr(resp.data[0], "url") and resp.data[0].url:
            img_resp = requests.get(resp.data[0].url, timeout=30)
            if img_resp.status_code == 200:
                return base64.b64encode(img_resp.content).decode(), None
        return None, "No image data in response"
    except Exception as e:
        return None, str(e)


def generate_images_batch(
    prompts: list[str],
    api_key: str,
    mode: str = "quick",
    width: int = 1024,
    height: int = 1024,
    ref_image_url: str | None = None,
    progress_callback=None,
) -> tuple[list[str | None], list[str]]:
    """Generate multiple images in parallel. Returns (list_of_base64, list_of_errors)."""
    results = [None] * len(prompts)
    errors = []

    def _gen(idx_prompt):
        idx, prompt = idx_prompt
        img, err = generate_single_image(prompt, api_key, mode, width, height, ref_image_url)
        return idx, img, err

    # Use fewer workers for Kontext (heavier model)
    max_w = 3 if mode == "shoe_plus" else 5
    with concurrent.futures.ThreadPoolExecutor(max_workers=max_w) as executor:
        futures = {executor.submit(_gen, (i, p)): i for i, p in enumerate(prompts)}
        done_count = 0
        for future in concurrent.futures.as_completed(futures):
            try:
                idx, img, err = future.result()
                results[idx] = img
                if err:
                    errors.append(f"Image {idx+1}: {err}")
            except Exception as e:
                errors.append(f"Thread error: {e}")
            done_count += 1
            if progress_callback:
                progress_callback(done_count, len(prompts))

    return results, errors


def extract_image_prompts(sections: dict) -> tuple[list[str], list[list[str]]]:
    """Extract image_prompt from static concepts and storyboard_prompts from video concepts.
    Returns (static_prompts: list[str], video_prompts: list[list[str]])."""
    static_prompts = []
    for concept in sections.get("static_concepts", []):
        # Look for image_prompt in markdown table: | image_prompt | <the prompt> |
        m = re.search(r"\|\s*image_prompt\s*\|\s*(.+?)\s*\|", concept, re.IGNORECASE)
        if m:
            static_prompts.append(m.group(1).strip())
        else:
            static_prompts.append("")

    video_prompts = []
    for concept in sections.get("video_concepts", []):
        # Look for storyboard_prompts in markdown table
        m = re.search(r"\|\s*storyboard_prompts?\s*\|\s*(.+?)\s*\|", concept, re.IGNORECASE)
        if m:
            raw = m.group(1).strip()
            # Split on "///" or "FRAME N:" patterns
            frames = re.split(r"\s*///\s*", raw)
            # Clean up FRAME N: prefix
            cleaned = []
            for f in frames:
                f = re.sub(r"^FRAME\s+\d+:\s*", "", f.strip(), flags=re.IGNORECASE)
                if f:
                    cleaned.append(f)
            video_prompts.append(cleaned[:3])  # max 3 frames
        else:
            video_prompts.append([])

    return static_prompts, video_prompts


# ──────────────────────────────────────────────
# STRATEGY PARSER & STRUCTURED RENDERER
# ──────────────────────────────────────────────

def parse_strategy(md: str) -> dict:
    """Split the raw markdown response into named sections."""
    sections = {
        "product_summary": "",
        "brand_lens": "",
        "static_concepts": [],
        "video_concepts": [],
        "priority_matrix": "",
        "creative_reference_map": "",
        "raw": md,
    }

    # Split on H1 headings (# HEADING)
    h1_pattern = re.compile(r"^# (.+)$", re.MULTILINE)
    h1_splits = h1_pattern.split(md)

    # h1_splits = [preamble, heading1, body1, heading2, body2, ...]
    current_key = None
    for i, part in enumerate(h1_splits):
        if i == 0:
            continue  # preamble before first heading
        if i % 2 == 1:
            # This is a heading name
            heading = part.strip().upper()
            if "PRODUCT SUMMARY" in heading:
                current_key = "product_summary"
            elif "BRAND LENS" in heading:
                current_key = "brand_lens"
            elif "STATIC" in heading and "CONCEPT" in heading:
                current_key = "static_concepts_raw"
            elif "VIDEO" in heading and "CONCEPT" in heading:
                current_key = "video_concepts_raw"
            elif "PRIORITY" in heading:
                current_key = "priority_matrix"
            elif "REFERENCE" in heading:
                current_key = "creative_reference_map"
            else:
                current_key = None
        else:
            # This is body content
            body = part.strip()
            if current_key in ("product_summary", "brand_lens", "priority_matrix", "creative_reference_map"):
                sections[current_key] = body
            elif current_key == "static_concepts_raw":
                # Split individual concepts on ## STATIC N:
                concepts = re.split(r"(?=^## STATIC \d+)", body, flags=re.MULTILINE)
                sections["static_concepts"] = [c.strip() for c in concepts if c.strip()]
            elif current_key == "video_concepts_raw":
                concepts = re.split(r"(?=^## VIDEO \d+)", body, flags=re.MULTILINE)
                sections["video_concepts"] = [c.strip() for c in concepts if c.strip()]

    return sections


def _extract_concept_title(concept_md: str) -> str:
    """Pull the concept title from a ## heading."""
    m = re.match(r"^##\s+(?:STATIC|VIDEO)\s+\d+:\s*[\"']?(.+?)[\"']?\s*$", concept_md, re.MULTILINE)
    return m.group(1).strip('"\'') if m else "Concept"


def render_strategy(sections: dict, static_images: list | None = None, video_images: list | None = None):
    """Render parsed strategy in structured tabs with styled cards and optional AI-generated visuals."""

    tab_overview, tab_static, tab_video, tab_priority = st.tabs([
        "Overview",
        f"Static Concepts ({len(sections['static_concepts'])})",
        f"Video Concepts ({len(sections['video_concepts'])})",
        "Priority & References",
    ])

    # ── Overview Tab ───────────────────────────
    with tab_overview:
        if sections["product_summary"]:
            st.markdown("### Product Summary")
            st.markdown(sections["product_summary"])
        if sections["brand_lens"]:
            st.markdown("---")
            st.markdown("### Brand Lens Applied")
            st.markdown(sections["brand_lens"])

    # ── Static Concepts Tab ────────────────────
    with tab_static:
        if not sections["static_concepts"]:
            st.info("No static concepts found in the response.")
        for i, concept in enumerate(sections["static_concepts"]):
            title = _extract_concept_title(concept)
            # Remove the ## heading line from the body
            body = re.sub(r"^##.+\n?", "", concept, count=1).strip()

            st.markdown(
                f'<div class="concept-card"><strong>{i+1}. {title}</strong></div>',
                unsafe_allow_html=True,
            )

            # Show generated image if available
            if static_images and i < len(static_images) and static_images[i]:
                img_col, text_col = st.columns([1, 1.5])
                with img_col:
                    st.image(
                        base64.b64decode(static_images[i]),
                        caption=f"AI-generated visual — Concept {i+1}",
                        use_container_width=True,
                    )
                with text_col:
                    with st.expander(f"View full concept — {title}", expanded=(i < 3)):
                        st.markdown(body)
            else:
                with st.expander(f"View full concept — {title}", expanded=(i < 3)):
                    st.markdown(body)

    # ── Video Concepts Tab ─────────────────────
    with tab_video:
        if not sections["video_concepts"]:
            st.info("No video concepts found in the response.")
        for i, concept in enumerate(sections["video_concepts"]):
            title = _extract_concept_title(concept)
            body = re.sub(r"^##.+\n?", "", concept, count=1).strip()

            st.markdown(
                f'<div class="concept-card video"><strong>{i+1}. {title}</strong></div>',
                unsafe_allow_html=True,
            )

            # Show storyboard frames if available
            if video_images and i < len(video_images) and video_images[i]:
                frames = video_images[i]
                if frames:
                    st.markdown("**Storyboard**")
                    frame_cols = st.columns(len(frames))
                    for fi, frame_b64 in enumerate(frames):
                        if frame_b64:
                            with frame_cols[fi]:
                                st.image(
                                    base64.b64decode(frame_b64),
                                    caption=f"Frame {fi+1}",
                                    use_container_width=True,
                                )

            with st.expander(f"View full concept — {title}", expanded=(i < 2)):
                st.markdown(body)

    # ── Priority & References Tab ──────────────
    with tab_priority:
        if sections["priority_matrix"]:
            st.markdown("### Priority Matrix")
            st.markdown(sections["priority_matrix"])
        if sections["creative_reference_map"]:
            st.markdown("---")
            st.markdown("### Creative Reference Map")
            st.markdown(sections["creative_reference_map"])
        if not sections["priority_matrix"] and not sections["creative_reference_map"]:
            st.info("No priority data found in the response.")


# ──────────────────────────────────────────────
# SESSION STATE DEFAULTS
# ──────────────────────────────────────────────
for key, default in {
    "strategy": None,
    "product_summary": None,
    "product_raw": None,
    "static_images": None,
    "video_images": None,
    "history": [],
}.items():
    if key not in st.session_state:
        st.session_state[key] = default


# ──────────────────────────────────────────────
# SIDEBAR
# ──────────────────────────────────────────────
with st.sidebar:
    st.markdown(
        "<h2 style='font-family:Bebas Neue,sans-serif;color:#EDD8C3;letter-spacing:3px;margin-bottom:0'>NEEMAN'S</h2>",
        unsafe_allow_html=True,
    )
    st.markdown("### Configuration")

    api_key = st.text_input(
        "Anthropic API Key",
        type="password",
        placeholder="sk-ant-api03-...",
        help="[Get a key](https://console.anthropic.com/)",
    )

    model = st.selectbox(
        "Model",
        [
            "claude-sonnet-4-20250514",
            "claude-opus-4-20250514",
            "claude-haiku-4-20250514",
        ],
        help="Sonnet → best quality / cost. Haiku → fastest.",
    )

    st.divider()
    st.markdown("### Visual Generation")
    visual_mode = st.radio(
        "Visual mode",
        options=["off", "quick", "shoe_plus"],
        format_func=lambda x: {
            "off": "No visuals (free)",
            "quick": "Quick visuals — Flux Schnell (~$0.08)",
            "shoe_plus": "Shoe+ Scenes — Flux Kontext (~$0.65)",
        }[x],
        index=0,
        help="Quick = generic AI visuals. Shoe+ = uses actual product shoe in each scene.",
    )
    visuals_on = visual_mode != "off"
    together_key = ""
    if visuals_on:
        together_key = st.text_input(
            "Together AI Key",
            type="password",
            placeholder="tok-...",
            help="[Get a free key](https://api.together.ai/)",
        )

    st.divider()
    st.markdown("### Email Delivery")
    email_on = st.toggle("Enable email delivery")
    smtp_email = smtp_pass = deliver_to = ""
    if email_on:
        smtp_email = st.text_input("Gmail address", placeholder="you@gmail.com")
        smtp_pass = st.text_input(
            "App password",
            type="password",
            help="Create at myaccount.google.com → Security → App passwords",
        )
        deliver_to = st.text_input(
            "Send results to",
            placeholder="team@neemans.com",
        )

    st.divider()
    brand_ctx = load_brand_context()
    if brand_ctx:
        st.success(f"Brand context loaded ({len(brand_ctx):,} chars)")
    else:
        st.error(
            "brand_context.md not found — place it next to app.py"
        )

    st.divider()

    # History
    if st.session_state["history"]:
        st.markdown("### History")
        for i, h in enumerate(reversed(st.session_state["history"][-5:])):
            st.caption(f"{h['time']}  —  {h['title']}")

    st.markdown("---")
    st.caption("Neeman's Creative Strategy Agent v1.0")
    st.caption("Built with Claude API + Streamlit")


# ──────────────────────────────────────────────
# MAIN CONTENT
# ──────────────────────────────────────────────
st.markdown(
    '<p class="main-header">Neeman\'s Creative Strategy Agent</p>',
    unsafe_allow_html=True,
)
st.markdown(
    '<p class="sub-header">Paste a product URL &rarr; get 10 static + 5 video ad concepts, brand-aligned and production-ready.</p>',
    unsafe_allow_html=True,
)

# URL input row
col_url, col_btn = st.columns([5, 1])
with col_url:
    product_url = st.text_input(
        "Product URL",
        placeholder="https://neemans.com/products/begin-walk-all-day-for-men-ivory-green",
        label_visibility="collapsed",
    )
with col_btn:
    go = st.button("Generate", type="primary", use_container_width=True)


# ──────────────────────────────────────────────
# GENERATE FLOW
# ──────────────────────────────────────────────
if go:
    # Guard-rails
    if not api_key:
        st.error("Enter your Anthropic API key in the sidebar.")
        st.stop()
    if not product_url or "neemans.com" not in product_url:
        st.error("Enter a valid neemans.com product URL.")
        st.stop()
    if not brand_ctx:
        st.error("Brand context file missing.")
        st.stop()

    # ── Step 1: Scrape ─────────────────────────
    with st.status("Scraping product data...", expanded=True) as status:
        product, err = scrape_product(product_url)
        if err:
            st.error(err)
            st.stop()
        summ = product_summary(product)
        status.update(label="Product scraped", state="complete")

    # ── Product preview card ───────────────────
    st.markdown("---")
    with st.container():
        img_col, info_col = st.columns([1, 2.4])
        with img_col:
            if summ["images"]:
                st.image(summ["images"][0], use_container_width=True)
        with info_col:
            st.markdown(f"### {summ['title']}")
            price_str = summ["price"]
            if summ["compare"]:
                price_str += f"  ~~{summ['compare']}~~"
            st.markdown(price_str)

            mcol1, mcol2, mcol3 = st.columns(3)
            mcol1.metric("Type", summ["type"] or "—")
            mcol2.metric("Sizes", f"{len(summ['sizes'])} available")
            mcol3.metric("Published", summ["published"] or "—")

            tags_html = " ".join(
                f'<span class="brand-tag">{t}</span>'
                for t in summ["tags"][:12]
            )
            st.markdown(tags_html, unsafe_allow_html=True)

    # ── Step 2: Generate ───────────────────────
    st.markdown("---")
    st.subheader("Creative Strategy")

    response_box = st.empty()
    full = ""
    with st.spinner("Generating creative strategy via Claude API..."):
        try:
            for chunk in stream_strategy(
                product, brand_ctx, api_key, model, product_url
            ):
                full += chunk
                response_box.markdown(full + " **|**")
        except Exception as e:
            st.error(f"API error: {e}")
            st.stop()

    # Replace streaming box with structured view
    response_box.empty()

    # Save to session
    st.session_state["strategy"] = full
    st.session_state["product_summary"] = summ
    st.session_state["product_raw"] = product
    st.session_state["history"].append(
        {
            "title": summ["title"],
            "time": datetime.now().strftime("%H:%M"),
            "url": product_url,
        }
    )

    st.success("Creative strategy generated!")

    # Parse and optionally generate visuals
    sections = parse_strategy(full)
    static_imgs = None
    video_imgs = None

    if visuals_on and together_key:
        static_prompts, video_prompt_lists = extract_image_prompts(sections)

        # For Shoe+ mode: get the product image URL and adapt prompts
        ref_image_url = None
        if visual_mode == "shoe_plus" and summ["images"]:
            ref_image_url = summ["images"][0]
            # Adapt prompts to work as image-editing instructions
            static_prompts = [
                f"Place this shoe in the following scene: {p}" if p else ""
                for p in static_prompts
            ]
            video_prompt_lists = [
                [f"Place this shoe in the following scene: {fp}" if fp else "" for fp in frames]
                for frames in video_prompt_lists
            ]

        mode_label = VISUAL_MODELS[visual_mode]["label"]

        # Gather all prompts for batch generation
        all_prompts = []
        prompt_map = []  # ("static", idx) or ("video", concept_idx, frame_idx)

        for idx, sp in enumerate(static_prompts):
            if sp:
                prompt_map.append(("static", idx))
                all_prompts.append(sp)

        for ci, frames in enumerate(video_prompt_lists):
            for fi, fp in enumerate(frames):
                if fp:
                    prompt_map.append(("video", ci, fi))
                    all_prompts.append(fp)

        if all_prompts:
            st.info(f"Found {len(all_prompts)} image prompts. Generating with **{mode_label}**...")
            progress_bar = st.progress(0, text=f"Generating {len(all_prompts)} visuals ({mode_label})...")

            def _update_progress(done, total):
                progress_bar.progress(done / total, text=f"Generating visuals ({mode_label})... {done}/{total}")

            raw_images, gen_errors = generate_images_batch(
                all_prompts,
                together_key,
                mode=visual_mode,
                ref_image_url=ref_image_url,
                progress_callback=_update_progress,
            )
            progress_bar.empty()

            # Map results back
            static_imgs = [None] * len(static_prompts)
            video_imgs = [[] for _ in video_prompt_lists]

            for i, mapping in enumerate(prompt_map):
                if mapping[0] == "static":
                    static_imgs[mapping[1]] = raw_images[i]
                elif mapping[0] == "video":
                    ci, fi = mapping[1], mapping[2]
                    while len(video_imgs[ci]) <= fi:
                        video_imgs[ci].append(None)
                    video_imgs[ci][fi] = raw_images[i]

            img_count = sum(1 for r in raw_images if r)
            if img_count > 0:
                st.success(f"Generated {img_count}/{len(all_prompts)} visuals with {mode_label}!")
            else:
                st.error("No visuals generated.")
            if gen_errors:
                with st.expander(f"⚠️ {len(gen_errors)} image generation issues"):
                    for e in gen_errors:
                        st.caption(e)
        else:
            st.warning("No image prompts found in the strategy output. Claude may not have included the image_prompt fields.")
    elif visuals_on and not together_key:
        st.warning("Enable visuals is on but no Together AI key entered.")

    # Save images to session state
    st.session_state["static_images"] = static_imgs
    st.session_state["video_images"] = video_imgs

    # Render structured output with visuals
    render_strategy(sections, static_images=static_imgs, video_images=video_imgs)

    # Raw markdown toggle
    if st.toggle("View raw markdown", value=False, key="raw_toggle_gen"):
        st.markdown(full)

    # ── Actions row ────────────────────────────
    st.markdown("---")
    act1, act2, act3 = st.columns(3)

    with act1:
        st.download_button(
            "Download .md",
            data=full,
            file_name=f"creative_strategy_{summ['handle']}.md",
            mime="text/markdown",
            use_container_width=True,
        )

    with act2:
        try:
            docx_buf = make_docx(summ["title"], full)
            st.download_button(
                "Download .docx",
                data=docx_buf,
                file_name=f"creative_strategy_{summ['handle']}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        except ImportError:
            st.info("Install `python-docx` for DOCX export.")

    with act3:
        if email_on and smtp_email and smtp_pass and deliver_to:
            if st.button("Send via Email", use_container_width=True):
                ok, err = send_email(
                    deliver_to,
                    f"Creative Strategy: {summ['title']}",
                    full,
                    {"email": smtp_email, "password": smtp_pass},
                )
                if ok:
                    st.success(f"Sent to {deliver_to}")
                else:
                    st.error(f"Email failed: {err}")
        else:
            st.button(
                "Email (configure in sidebar)",
                disabled=True,
                use_container_width=True,
            )


# ──────────────────────────────────────────────
# SHOW LAST RESULT (on re-run / no new generation)
# ──────────────────────────────────────────────
elif st.session_state["strategy"]:
    summ = st.session_state["product_summary"]
    full = st.session_state["strategy"]

    st.markdown("---")
    st.caption(f"Showing last generated strategy for **{summ['title']}**")

    # Structured view with cached images
    sections = parse_strategy(full)
    render_strategy(
        sections,
        static_images=st.session_state.get("static_images"),
        video_images=st.session_state.get("video_images"),
    )

    # Raw toggle
    if st.toggle("View raw markdown", value=False, key="raw_toggle_last"):
        st.markdown(full)

    st.markdown("---")
    a1, a2, _ = st.columns(3)
    with a1:
        st.download_button(
            "Download .md",
            data=full,
            file_name=f"creative_strategy_{summ['handle']}.md",
            mime="text/markdown",
            use_container_width=True,
        )
    with a2:
        try:
            docx_buf = make_docx(summ["title"], full)
            st.download_button(
                "Download .docx",
                data=docx_buf,
                file_name=f"creative_strategy_{summ['handle']}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        except ImportError:
            pass

else:
    # Landing state — show example cards
    st.markdown("---")
    st.markdown("#### How it works")
    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown(
            """
            <div class="stat-box">
                <div class="stat-value">1</div>
                <div class="stat-label">Paste product URL</div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    with c2:
        st.markdown(
            """
            <div class="stat-box">
                <div class="stat-value">2</div>
                <div class="stat-label">AI generates strategy</div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    with c3:
        st.markdown(
            """
            <div class="stat-box">
                <div class="stat-value">3</div>
                <div class="stat-label">Download or email</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    st.markdown("")
    st.markdown("**What you get:**")
    st.markdown(
        """
- 10 static ad concepts (hook, visual, copy, format, reasoning)
- 5 video ad concepts (scene-by-scene breakdown, scripts)
- AI-generated hero visuals for each static concept
- 3-frame storyboards for each video concept
- Priority matrix with top 3 must-produce picks
- Every concept aligned to Neeman's brand guidelines
- Download as `.md` or `.docx`, or send via email
"""
    )
